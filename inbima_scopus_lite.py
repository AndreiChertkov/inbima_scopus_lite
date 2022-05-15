import csv
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Mm, Cm, Pt
import os
import sys


from excel import Excel
from journals_description import description

sys.path.append('./data')
from journals_map import journals_map
from papers_del import papers_del


FOLD_AUTHORS = './data'
FOLD_RESULT = './result'
PATH_JOURNALS = './data/journals.xlsx'


def export(papers, fpath):
    document = docx.Document()

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    i = 0
    for q in [1, 2, 0]:
        if q == 1:
            text = 'Publications in Q1-rated journals\n'
        elif q == 2:
            text = 'Publications in Q2-rated journals\n'
        else:
            text = 'Other publications\n'
        document.add_heading(text, level=3)

        j = 0
        for item in papers.values():
            if item['q'] != q: continue
            i += 1
            j += 1

            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            res = p.add_run(f'[{i}] ')

            res = p.add_run(item["title"] + '. ')
            res.bold = True

            text = ', '.join(item["authors"])
            text = text.strip()
            if text[-1] != '.':
                text += '.'
            text += ' '
            res = p.add_run(text)

            res = p.add_run(f'{item["journal"]}. {item["year"]}. ')

            if item["volume"]:
                text = f'V. {item["volume"]}'
                if item["number"]:
                    text += f' № {item["number"]}'
                text += '.'
                res = p.add_run(text)

        print(f'... Saved Q{q} papers [{j:-3d}]')

    document.save(fpath)


def inbima_scopus_lite(year_min=None):
    journals = Excel(PATH_JOURNALS, description, _log).load('journals')
    print(f'>>> Loaded all journals [{len(journals)}]')

    papers_all = {}
    for uid in _get_uids():
        print(f'... Parsing {uid}')
        fpath = os.path.join(FOLD_AUTHORS, uid + '.csv')
        papers = parse(fpath, journals, year_min)
        print(f'>>> Parsed  {uid} [{len(papers)}]')

        print(f'... Saving  {uid}')
        fname = f'result_{uid}'
        if year_min:
            fname += f'_{year_min}'
        fpath = os.path.join(FOLD_RESULT, fname + '.docx')
        export(papers, fpath)
        print(f'>>> Saved   {uid} [{len(papers)}]')

        papers_all.update(papers)

    print(f'... Saving  all')
    fname = f'result'
    if year_min:
        fname += f'_{year_min}'
    fpath = os.path.join(FOLD_RESULT, fname + '.docx')
    export(papers_all, fpath)
    print(f'>>> Saved   all [{len(papers_all)}]')


def parse(fpath, journals={}, year_min=None):
    papers = {}

    with open(fpath, newline='') as f:
        for i, row in enumerate(csv.reader(f, delimiter=',')):
            if i == 0:
                continue

            if year_min and int(row[3]) < year_min:
                continue

            title = row[2]
            if title in papers_del:
                continue

            journal = row[4]
            if journal in journals_map:
                journal = journals_map[journal]
            journal = journals.get(journal, {})

            q = 0
            if not journal:
                print(f'WRN Can not find journal')
                print(f'Journal : "{row[4]}"')
                print(f'Title   : "{title}"\n')
                print(f'')
            else:
                if journal.get('q1'):
                    q = 1
                elif journal.get('q2'):
                    q = 2

            papers[title] = {
                'title': title,
                'authors': row[0].split(', '),
                'year': row[3],
                'journal': row[4],
                'volume': row[5],
                'number': row[6],
                'page1': row[8],
                'page2': row[9],
                'kind': row[49] if len(row) > 48 else None,
                'q': q,
            }

    return papers


def _get_uids():
    uids = []
    for fpath in os.listdir(FOLD_AUTHORS):
        if fpath.endswith('.csv'):
            uids.append(fpath.split('.csv')[0])
    return uids


def _log(text, kind='res'):
    print(kind.upper() + ' ' + text)


if __name__ == '__main__':
    year_min = int(sys.argv[1]) if len(sys.argv) > 1 else None

    inbima_scopus_lite(year_min)
